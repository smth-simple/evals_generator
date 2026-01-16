#!/usr/bin/env python3
import os
import sys
import json
import logging
import asyncio
import argparse
import subprocess
import re
from datetime import datetime
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from litellm import completion

# Configure Logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --- FIGURE DESCRIPTION MAPPING (Descriptions ONLY, Numbering is Dynamic) ---
FIGURE_DESCRIPTIONS = {
    "viz_00": "Global Performance Summary: Net Win Rate & Major Error Deltas.",
    "viz_09": "Truthfulness Heatmap: Frequency of Hallucinations by Region.",
    "viz_10": "Instruction Following Heatmap: Constraint Failures by Region.",
    "viz_13": "Capability Analysis: Win Rates across Generative vs. Reference Tasks.",
    "viz_14": "Complexity Analysis: Performance on High-Difficulty LMArena Categories.",
    "viz_15": "Polish & Formatting: Analysis of Minor Error Differentials.",
    "viz_16": "Quality Composition: Contribution of IF, TF, and Loc to Overall Score.",
    "viz_17": "User Sentiment: Distribution of Preference Intensity (Likert 1-7).",
    "viz_18": "Style Divergence: Response Length Tendencies (Concise vs. Verbose).",
    "viz_19": "Length Disparity: Top 8 Languages with Extreme Length Differences.",
    "viz_20": "Preference Disparity: Top 8 Languages with Highest User Disagreement."
}

class DatasetProcessor:
    def __init__(self, dataset_path):
        self.df = pd.read_csv(dataset_path)
        
        # --- 1. COMPUTE MAJOR & MINOR ISSUES ---
        def check_major_issue(row):
            metrics = ['instruction_following', 'truthfulness', 'localization', 
                       'writing_style_tone', 'harmlessness']
            for m in metrics:
                if row.get(m, 3) == 1: return 1
            rl = row.get('response_length', 0)
            if rl == -2 or rl == 2: return 1
            return 0

        self.df['is_major_issue'] = self.df.apply(check_major_issue, axis=1)
        self.region_map = {
            "ja_JP": "East Asia", "ko_KR": "East Asia", "zh_CN": "East Asia", "zh_HK": "East Asia", "zh_TW": "East Asia",
            "hi_IN": "South Asia",
            "id_ID": "Southeast Asia", "ms_MY": "Southeast Asia", "th_TH": "Southeast Asia", "vi_VN": "Southeast Asia",
            "ar_AE": "MENA", "ar_EG": "MENA", "ar_SA": "MENA", "he_IL": "MENA", "tr_TR": "MENA",
            "fr_FR": "Western Europe", "fr_BE": "Western Europe", "fr_CH": "Western Europe",
            "de_DE": "Western Europe", "de_AT": "Western Europe", "de_CH": "Western Europe",
            "it_IT": "Western Europe", "it_CH": "Western Europe", "es_ES": "Western Europe",
            "pt_PT": "Western Europe", "nl_NL": "Western Europe", "nl_BE": "Western Europe",
            "da_DK": "Nordics", "fi_FI": "Nordics", "no_NO": "Nordics", "sv_SE": "Nordics",
            "pl_PL": "Eastern Europe", "ru_RU": "Eastern Europe", "uk_UA": "Eastern Europe",
            "es_CL": "Latin America", "es_MX": "Latin America", "es_US": "Latin America", "pt_BR": "Latin America",
            "en_US": "North America", "fr_CA": "North America",
        }
        self.df['region'] = self.df['language_code'].map(self.region_map).fillna('Other')

    def compute_stats(self):
        """Computes the 'Quantitative Backbone'."""
        def fmt(x): return f"{x:.1f}%"
        
        unique_tasks = self.df.drop_duplicates(subset=['task_id']).copy()
        total = len(unique_tasks)
        if total > 0:
            wins_a = len(unique_tasks[unique_tasks['likert'] < 4])
            ties = len(unique_tasks[unique_tasks['likert'] == 4])
            wins_b = len(unique_tasks[unique_tasks['likert'] > 4])
            win_summary = pd.DataFrame({
                'Outcome': ['Base Model Wins (1-3)', 'Ties (4)', 'Test Model Wins (5-7)'],
                'Count': [wins_a, ties, wins_b],
                'Percentage': [fmt(wins_a/total*100), fmt(ties/total*100), fmt(wins_b/total*100)]
            }).to_markdown(index=False)
        else: win_summary = "No data."

        unique_tasks['win_a'] = unique_tasks['likert'] < 4
        unique_tasks['win_b'] = unique_tasks['likert'] > 4
        
        def build_win_table(group_col):
            df_ = unique_tasks.groupby(group_col)[['win_a', 'win_b']].mean() * 100
            df_['Net Win Rate (Test-Base)'] = df_['win_b'] - df_['win_a']
            return df_.sort_values('Net Win Rate (Test-Base)', ascending=False).reset_index().to_markdown(floatfmt=".1f")

        win_lang = build_win_table('language_code')
        win_region = build_win_table('region')
        win_prompt_cat = build_win_table('prompt_category') if 'prompt_category' in unique_tasks.columns else "N/A"

        metrics = ['instruction_following', 'truthfulness', 'localization', 'writing_style_tone', 'harmlessness']
        def get_error_rates(df_subset):
            stats = {}
            for m in metrics:
                stats[f"{m} (Major)"] = (df_subset[m] == 1).mean() * 100
            return pd.Series(stats)

        try:
            overall_errors = pd.DataFrame({
                'Base Model': get_error_rates(self.df[self.df['model_name']=='model_a']), 
                'Test Model': get_error_rates(self.df[self.df['model_name']=='model_b'])
            }).reset_index().to_markdown(floatfmt=".1f")
        except: overall_errors = "N/A"

        def build_error_pivot(idx_col):
            piv = self.df.pivot_table(index=idx_col, columns='model_name', values='is_major_issue', aggfunc='mean') * 100
            if 'model_a' in piv.columns and 'model_b' in piv.columns:
                piv['Diff (Test-Base)'] = piv['model_b'] - piv['model_a']
            return piv.sort_values('model_b', ascending=False).reset_index().to_markdown(floatfmt=".1f")

        lang_error_table = build_error_pivot('language_code')
        region_error_table = build_error_pivot('region')
        prompt_cat_table = build_error_pivot('prompt_category') if 'prompt_category' in self.df.columns else "N/A"
        occ_table = build_error_pivot('occupational_category') if 'occupational_category' in self.df.columns else "N/A"

        # LMArena
        lm_summary = pd.DataFrame(index=self.df['language_code'].unique())
        def add_col(name, filter_col, val, model):
            mask = (self.df[filter_col] == val) & (self.df['model_name'] == model)
            if mask.any(): lm_summary[name] = self.df[mask].groupby('language_code')['is_major_issue'].mean() * 100
        
        if 'hard_prompts' in self.df.columns:
            add_col('Hard(Base)', 'hard_prompts', 'Hard Prompt', 'model_a')
            add_col('Hard(Test)', 'hard_prompts', 'Hard Prompt', 'model_b')
        lm_arena_table = lm_summary.dropna(how='all').reset_index().to_markdown(floatfmt=".1f")

        # Subcategories
        subcat_cols = ['loc_subcategory', 'if_subcategory', 'tf_subcategory', 'rl_subcategory']
        cat_map = {'loc_subcategory': 'Loc','if_subcategory': 'IF','tf_subcategory': 'TF','rl_subcategory': 'RL'}
        melted = self.df[self.df['is_major_issue']==1].melt(id_vars=['model_name', 'language_code'], value_vars=subcat_cols, var_name='source', value_name='subcat').dropna()
        melted['cat'] = melted['source'].map(cat_map)
        
        grouped = melted.groupby(['model_name', 'language_code', 'cat', 'subcat']).size().reset_index(name='count')
        totals = self.df.groupby(['model_name', 'language_code']).size().reset_index(name='total')
        merged = pd.merge(grouped, totals, on=['model_name', 'language_code'])
        merged['rate'] = (merged['count'] / merged['total']) * 100
        subcat_table = merged.sort_values(['model_name', 'language_code', 'count'], ascending=[True, True, False]).groupby(['model_name', 'language_code']).head(5).to_markdown(index=False, floatfmt=".2f")

        return {
            "win_summary": win_summary, "win_lang": win_lang, "win_region": win_region, "win_prompt_cat": win_prompt_cat,
            "overall_errors": overall_errors, "lang_error_table": lang_error_table, "region_error_table": region_error_table,
            "prompt_cat_table": prompt_cat_table, "occ_table": occ_table, "lm_arena_table": lm_arena_table, "subcat_table": subcat_table
        }

    # --- SAMPLING HELPERS ---
    def get_lean_failure_samples(self, n=10, filter_criteria=None, model_name=None):
        df_failures = self.df[self.df['is_major_issue'] == 1].copy()
        if model_name: df_failures = df_failures[df_failures['model_name'] == model_name]
        if filter_criteria:
            if 'languages' in filter_criteria: df_failures = df_failures[df_failures['language_code'].isin(filter_criteria['languages'])]
            if 'regions' in filter_criteria: df_failures = df_failures[df_failures['region'].isin(filter_criteria['regions'])]
            if 'is_lmarena' in filter_criteria:
                col = 'lm_arena_category' if 'lm_arena_category' in df_failures.columns else 'category'
                if col in df_failures.columns: df_failures = df_failures[df_failures[col].notna()]
        
        if len(df_failures) > n: df_failures = df_failures.sample(n)
        elif len(df_failures) == 0: return [] 

        lean_samples = []
        for _, row in df_failures.iterrows():
            lean_samples.append({
                "task_id": row.get('task_id', 'N/A'), "model": row.get('model_name'), "lang": row.get('language_code'), "subcat": row.get('subcategory', 'N/A'),
                "justification": row.get('justification', 'N/A'), "prompt": str(row.get('prompt', ''))[:300]
            })
        return lean_samples

    def get_stratified_region_samples(self, region_name, top_n_subcats=3, examples_per_subcat=5, model_name=None):
        mask = (self.df['region'] == region_name) & (self.df['is_major_issue'] == 1)
        if model_name: mask = mask & (self.df['model_name'] == model_name)
        region_df = self.df[mask].copy()
        if region_df.empty: return []

        def get_primary_subcat(row):
            for col in ['loc_subcategory', 'if_subcategory', 'tf_subcategory', 'rl_subcategory']:
                val = row.get(col)
                if pd.notna(val):
                    s_val = str(val).strip()
                    if s_val.lower() not in ['nan', 'none', 'no issues', 'no issue', 'n/a']: return s_val
            return 'Uncategorized' 

        if 'subcategory' not in region_df.columns:
            region_df['subcategory'] = region_df.apply(get_primary_subcat, axis=1)

        samples = []
        for lang in region_df['language_code'].unique():
            lang_df = region_df[region_df['language_code'] == lang]
            valid_subcats = lang_df[~lang_df['subcategory'].isin(['Uncategorized', 'No Issues', 'No Issue'])]
            top_subcats = valid_subcats['subcategory'].value_counts().head(top_n_subcats).index.tolist()
            for subcat in top_subcats:
                picks = subcat_df = lang_df[lang_df['subcategory'] == subcat].sample(min(len(lang_df[lang_df['subcategory'] == subcat]), examples_per_subcat))
                for _, row in picks.iterrows():
                    metrics = ['instruction_following', 'truthfulness', 'localization', 'response_length']
                    active_failures = [m for m in metrics if row.get(m, 0) == 1]
                    samples.append({
                        "task_id": row.get('task_id', 'N/A'),
                        "language": lang, "region": region_name, "model": model_name, "failure_type": subcat,
                        "failed_metrics": active_failures,
                        "justification": row.get('justification', 'N/A'), 
                        "prompt": str(row.get('prompt', ''))[:300]
                    })
        return samples

    def get_interesting_languages(self, top_n=5):
        unique_tasks = self.df.drop_duplicates(subset=['task_id']).copy()
        unique_tasks['win_b'] = unique_tasks['likert'] > 4
        win_rates = unique_tasks.groupby('language_code')['win_b'].mean()
        return win_rates.sort_values(ascending=False).head(top_n).index.tolist()

    def get_head_to_head_samples(self, language_code, n=8):
        lang_df = self.df[(self.df['language_code'] == language_code)].drop_duplicates(subset=['task_id'])
        a_wins = lang_df[lang_df['likert'] < 4]
        b_wins = lang_df[lang_df['likert'] > 4]
        samples = []
        for _, row in a_wins.sample(min(len(a_wins), n)).iterrows():
            samples.append({
                "type": "BASE MODEL WIN", "task_id": row.get('task_id'), "likert": row.get('likert'),
                "justification": row.get('justification', 'N/A'), "prompt": str(row.get('prompt', ''))[:300]
            })
        for _, row in b_wins.sample(min(len(b_wins), n)).iterrows():
            samples.append({
                "type": "TEST MODEL WIN", "task_id": row.get('task_id'), "likert": row.get('likert'),
                "justification": row.get('justification', 'N/A'), "prompt": str(row.get('prompt', ''))[:300]
            })
        return samples
    
    def get_lang_stats(self, lang):
        a_err = self.df[(self.df['language_code']==lang) & (self.df['model_name']=='model_a')]['is_major_issue'].mean() * 100
        b_err = self.df[(self.df['language_code']==lang) & (self.df['model_name']=='model_b')]['is_major_issue'].mean() * 100
        tasks = self.df[self.df['language_code']==lang].drop_duplicates(subset=['task_id'])
        a_win = (tasks['likert'] < 4).mean() * 100
        b_win = (tasks['likert'] > 4).mean() * 100
        return {"a_error": f"{a_err:.1f}", "b_error": f"{b_err:.1f}", "a_win": f"{a_win:.1f}", "b_win": f"{b_win:.1f}"}

class ReportGeneratorV2:
    def __init__(self, config_path):
        with open(config_path) as f: self.config = json.load(f)
        with open(self.config['prompts_file']) as f: self.prompts_config = json.load(f)
        with open(self.config['models_file']) as f: self.models_config = json.load(f)
        self.token_usage = {'total_tokens': 0, 'api_calls': 0}

    # --- STATE MANAGEMENT ---
    def _save_state(self, filename, data):
        path = Path(self.config['report_output_dir']) / filename
        path.parent.mkdir(parents=True, exist_ok=True)
        with open(path, 'w') as f: json.dump(data, f, default=str)
        logger.info(f"  💾 State saved: {filename}")

    def _load_state(self, filename):
        path = Path(self.config['report_output_dir']) / filename
        if not path.exists():
            sys.exit(1)
        with open(path, 'r') as f: return json.load(f)

    def _get_model_config(self, step_name):
        return self.models_config.get('steps', {}).get(step_name, self.models_config['default'])

    async def _call_llm(self, step, prompt, context):
        cfg = self._get_model_config(step)
        ctx_str = json.dumps(context, default=str)
        if len(ctx_str) > 300000: ctx_str = ctx_str[:300000] + "... [TRUNCATED]"
        
        system_prompt = "You are a professional data analyst. Do not use 'Interpretation:' labels."

        try:
            resp = completion(
                model=cfg['model'],
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"{prompt}\n\nCONTEXT:\n{ctx_str}"}
                ],
                temperature=cfg['temperature'],
                max_tokens=cfg.get('max_tokens', 4000)
            )
            content = resp.choices[0].message.content
            if hasattr(resp, 'usage'): self.token_usage['total_tokens'] += resp.usage.total_tokens
            return {'content': content}
        except Exception as e:
            logger.error(f"LLM Error in {step}: {e}")
            return {'content': f"[Error: {e}]"}

    # --- FORMATTING LOGIC ---
    def _apply_markdown_formatting(self, paragraph, text):
        """Applies bolding to a paragraph object."""
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                clean_text = part[2:-2]
                run = paragraph.add_run(clean_text)
                run.bold = True
            else:
                paragraph.add_run(part)

    def _render_markdown_table_to_docx(self, doc, md_string):
        """Converts Markdown table to Word Table with BOLD SUPPORT."""
        lines = [l.strip() for l in md_string.strip().split('\n') if l.strip()]
        if len(lines) < 2: return
        
        headers = [h.strip() for h in lines[0].split('|') if h.strip()]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Header Row
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]
            p.clear() 
            self._apply_markdown_formatting(p, h)
            for run in p.runs: run.font.bold = True 
        
        # Data Rows
        for line in lines[2:]:
            if '---' in line: continue
            row_data = [d.strip() for d in line.split('|') if d.strip()]
            if len(row_data) != len(headers): continue
            
            row_cells = table.add_row().cells
            for i, d in enumerate(row_data):
                p = row_cells[i].paragraphs[0]
                p.clear()
                self._apply_markdown_formatting(p, d)

    def _get_clean_caption(self, filename):
        """Generates a clean caption from filename using fuzzy matching on FIGURE_DESCRIPTIONS."""
        # 1. Try exact match key (e.g. 'viz_00')
        for key, desc in FIGURE_DESCRIPTIONS.items():
            if key in filename:
                return desc
        
        # 2. Fallback: Clean the filename
        clean = filename.replace('.png', '').replace('viz_', '').replace('_', ' ').title()
        return f"{clean} (Analysis)"

    def _save_docx(self, report, output_path):
        doc = Document()
        
        # Title
        title = doc.add_heading(self.config.get('report_title', 'Rubrics LLM Report'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # --- PHASE 1: PRE-SCAN FOR FIGURES (Assign Numbers Sequentially) ---
        viz_map = {}
        if Path(self.config['viz_output_dir']).exists():
            for viz_path in sorted(Path(self.config['viz_output_dir']).glob("*.png")):
                viz_map[viz_path.name] = str(viz_path)

        # Scan all sections to find figure tags in order
        figure_order = []
        for section in report['sections']:
            # Find all [[FIGURE: ...]] tags
            tags = re.findall(r'\[\[FIGURE:\s*(.*?)\]\]', section['content'])
            for tag in tags:
                clean_tag = tag.strip()
                if clean_tag not in figure_order:
                    figure_order.append(clean_tag)
        
        # Map filename -> "Figure X"
        fig_num_map = {name: i+1 for i, name in enumerate(figure_order)}

        # --- PHASE 2: RENDER DOCUMENT WITH REPLACEMENTS ---
        for section in report['sections']:
            doc.add_heading(section['title'], level=1)
            
            # Split by Table OR Figure Tag
            parts = re.split(r'((\|.*\|[\r\n]+\|[-:| ]+\|[\r\n]+(?:\|.*\|[\r\n]*)+)|(\[\[FIGURE:.*?\]\]))', section['content'], flags=re.DOTALL)
            
            for part in parts:
                if not part: continue
                part = part.strip()
                if not part: continue

                # A: Table
                if part.startswith('|') and '---' in part:
                    self._render_markdown_table_to_docx(doc, part)
                
                # B: Figure Tag
                elif part.startswith('[[FIGURE:') and part.endswith(']]'):
                    filename = part.replace('[[FIGURE:', '').replace(']]', '').strip()
                    if filename in viz_map:
                        try:
                            doc.add_picture(viz_map[filename], width=Inches(6.0))
                            
                            # Dynamic Numbering
                            fig_num = fig_num_map.get(filename, "?")
                            desc = self._get_clean_caption(filename)
                            
                            cap = doc.add_paragraph(f"Figure {fig_num}. {desc}", style='Caption')
                            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            logger.error(f"Failed to add image {filename}: {e}")
                    else:
                        doc.add_paragraph(f"[Image Not Found: {filename}]", style="Quote")

                # C: Text (Perform replacements here!)
                else:
                    # Replace (viz_00...) with (Figure 1)
                    processed_text = part
                    for fname, fnum in fig_num_map.items():
                        # Replace full filename refs
                        processed_text = processed_text.replace(fname, f"Figure {fnum}")
                        # Replace lazy refs "viz 00"
                        short_name = fname.split('_')[0] + "_" + fname.split('_')[1] # viz_00
                        processed_text = processed_text.replace(short_name, f"Figure {fnum}")
                    
                    # Render processed text
                    for line in processed_text.split('\n'):
                        line = line.strip()
                        if not line: continue
                        if line.startswith('### '): doc.add_heading(line[4:], level=3)
                        elif line.startswith('## '): doc.add_heading(line[3:], level=2)
                        elif line.startswith('# '): doc.add_heading(line[2:], level=1)
                        elif line.startswith('* ') or line.startswith('- '):
                            self._apply_markdown_formatting(doc.add_paragraph(style='List Bullet'), line.lstrip('*- '))
                        else:
                            self._apply_markdown_formatting(doc.add_paragraph(style='Normal'), line)
            
            doc.add_page_break()
            
        doc.save(output_path)
        logger.info(f"  ✓ DOCX Saved: {output_path}")

    # --- STEPS ---
    async def step_understanding(self): return {'content': "DATA DICTIONARY: [Same as before]"}
    
    def step_visualize(self):
        logger.info("  - Running Visualization Scripts...")
        viz_dir = Path(self.config['viz_scripts_dir'])
        if not viz_dir.exists(): return
        env = os.environ.copy()
        env['DATASET_PATH'] = self.config['dataset_path']
        env['OUTPUT_DIR'] = self.config['viz_output_dir']
        Path(self.config['viz_output_dir']).mkdir(parents=True, exist_ok=True)
        for script in sorted(viz_dir.glob("*.py")):
            try: subprocess.run([sys.executable, str(script)], env=env, check=True)
            except: pass

    async def step_identifying(self, understanding):
        logger.info("  - Running Deep Stratified Analysis...")
        processor = DatasetProcessor(self.config['dataset_path'])
        stats = processor.compute_stats()
        
        global_stats_context = f"""# GLOBAL STATISTICS\n{stats['win_summary']}\n{stats['win_prompt_cat']}\n{stats['overall_errors']}\n{stats['prompt_cat_table']}\n{stats['occ_table']}"""
        regional_stats_context = f"""# REGIONAL STATISTICS\n{stats['win_region']}\n{stats['region_error_table']}\n{stats['win_lang']}\n{stats['lang_error_table']}\n{stats['subcat_table']}"""
        lm_stats_context = f"""# COMPLEXITY STATISTICS\n{stats['lm_arena_table']}"""

        full_analysis_output = [global_stats_context + "\n\n" + regional_stats_context]
        
        # Per-Model Loop
        for model in ['model_a', 'model_b']:
            model_label = "Base Model (model_a)" if model == 'model_a' else "Test Model (model_b)"
            model_insights = []
            
            # Global
            global_samples = processor.get_lean_failure_samples(n=15, model_name=model)
            model_insights.append(await self._run_analysis_pass(f"{model_label} - Global", "identifying_global", global_samples, f"Analyze global failures.\nRELEVANT DATA:\n{global_stats_context}"))

            # Regional
            m_df = processor.df[processor.df['model_name'] == model]
            if not m_df.empty:
                bad_regions = m_df[m_df['is_major_issue']==1]['region'].value_counts().head(3).index.tolist()
                for region in bad_regions:
                    if region == 'Other': continue
                    region_samples = processor.get_stratified_region_samples(region, top_n_subcats=3, examples_per_subcat=5, model_name=model)
                    if region_samples:
                        insight = await self._run_analysis_pass(f"{model_label} - {region}", "identifying_region_loop", region_samples, f"Focus on {region}.\nRELEVANT STATS:\n{regional_stats_context}")
                        model_insights.append(f"### {model_label}: {region} Analysis\n{insight}")
            
            # LMArena
            lm_samples = processor.get_lean_failure_samples(n=15, filter_criteria={'is_lmarena': True}, model_name=model)
            model_insights.append(await self._run_analysis_pass(f"{model_label} - LMArena", "identifying_lmarena", lm_samples, f"Analyze Hard Prompts.\nRELEVANT STATS:\n{lm_stats_context}"))
            
            full_analysis_output.append(f"# ANALYSIS OF {model_label}\n" + "\n\n".join(model_insights))

        # Comparative Loop
        target_langs = processor.get_interesting_languages(top_n=5)
        h2h_insights = []
        for lang in target_langs:
            lang_stats = processor.get_lang_stats(lang)
            h2h_samples = processor.get_head_to_head_samples(lang, n=8)
            prompt_key = "identifying_h2h"
            prompt_text = self.prompts_config.get(prompt_key, {}).get('prompt', "Compare.").format(language_name=lang, **lang_stats)
            insight = (await self._call_llm(step="identifying", prompt=prompt_text, context={'sample_data': h2h_samples}))['content']
            h2h_insights.append(f"### Head-to-Head: {lang}\n{insight}")

        full_analysis_output.append("# COMPARATIVE ANALYSIS\n" + "\n\n".join(h2h_insights))
        final_content = "\n\n".join(full_analysis_output)
        self._save_file("identifying_output.md", final_content)
        return {'content': final_content, 'raw_stats': stats, 'understanding_context': understanding['content']}

    async def _run_analysis_pass(self, pass_name, prompt_key, samples, extra_context):
        prompt_text = self.prompts_config.get(prompt_key, {}).get('prompt', f"Analyze {pass_name}.")
        return (await self._call_llm(step="identifying", prompt=prompt_text, context={'analysis_goal': extra_context, 'sample_data': samples}))['content']

    async def step_synthesis(self, identifying):
        logger.info("  - Running Narrative Synthesis...")
        id_content = identifying['content']
        viz_dir = Path(self.config['viz_output_dir'])
        csv_data_summary = []
        for f in sorted(viz_dir.glob("*.csv")):
             try:
                df = pd.read_csv(f)
                csv_data_summary.append(f"### {f.stem}\n{df.to_string(index=False)}")
             except: pass
        data_context_str = "\n\n".join(csv_data_summary)
        
        synth_config = self.prompts_config.get('synthesis_prompt', {})
        full_prompt = f"{synth_config.get('dictionary', '')}\n\n# QUALITATIVE\n{id_content}\n\n# QUANTITATIVE\n{data_context_str}\n\n{synth_config.get('prompt', '')}"

        resp = await self._call_llm(step="synthesis", prompt=full_prompt, context={})
        self._save_file("synthesis_output.md", resp['content'])
        return {'content': resp['content'], 'identifying_content': id_content}

    async def step_rollout(self, synthesis_data):
        logger.info("  - Generating Final Report Sections...")
        strategic_brief = synthesis_data['content']
        id_insights = synthesis_data['identifying_content']

        viz_dir = Path(self.config['viz_output_dir'])
        figure_list = [f"- {f.name}" for f in viz_dir.glob("*.png")]
        figure_list_str = "\n".join(figure_list)
        
        csv_tables = []
        for f in sorted(viz_dir.glob("*.csv")):
            try:
                df = pd.read_csv(f)
                if len(df) < 50: csv_tables.append(f"### Data Table: {f.stem}\n{df.to_markdown(index=False)}")
            except: pass
        csv_tables_str = "\n\n".join(csv_tables)

        sections = []
        for section_cfg in self.prompts_config['report_sections']:
            logger.info(f"    - {section_cfg['title']}...")
            context_payload = {
                'strategic_brief': strategic_brief,
                'identifying_insights': id_insights[:50000], 
                'available_figures': figure_list_str,
                'data_tables': csv_tables_str
            }
            resp = await self._call_llm(step="rollout", prompt=section_cfg['prompt'], context=context_payload)
            sections.append({"title": section_cfg['title'], "content": resp['content']})
            
        return {'sections': sections, 'variant': 1}
    
    async def step_combine_check(self, report):
        full_text = "\n".join([s['content'] for s in report['sections']])
        await self._call_llm(step="combine_check", prompt="Check", context={'text': full_text[:50000]})
        self._save_single_report_docx(report)

    def _save_file(self, filename, content):
        path = Path(self.config['report_output_dir']) / filename
        path.parent.mkdir(parents=True, exist_ok=True)
        with open(path, 'w', encoding='utf-8') as f: f.write(content)

    def _save_single_report_docx(self, report):
        output_dir = Path(self.config['report_output_dir'])
        filename = output_dir / f"report_final.docx"
        self._save_docx(report, filename)

    async def generate_report(self, mode='all'):
        logger.info(f"\n📝 STARTING REPORT GENERATION (Mode: {mode})")
        if mode in ['all', 'viz']: self.step_visualize()
        if mode == 'viz': return
        if mode in ['all', 'identifying']:
            u = await self.step_understanding()
            i = await self.step_identifying(u)
            self._save_state('state_identifying.json', i)
        else: i = self._load_state('state_identifying.json')
        if mode == 'identifying': return

        if mode in ['all', 'synthesis']:
            s = await self.step_synthesis(i)
            self._save_state('state_synthesis.json', s)
        else: s = self._load_state('state_synthesis.json')
        if mode == 'synthesis': return

        if mode in ['all', 'rollout', 'rollout_check']:
            r = await self.step_rollout(s)
            self._save_state('state_rollout.json', r)
        elif mode == 'check': r = self._load_state('state_rollout.json')
        if mode == 'rollout': return

        if mode in ['all', 'check', 'rollout_check']: await self.step_combine_check(r)

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--config', default='evals_report_config_v2.json')
    parser.add_argument('--step', default='all')
    args = parser.parse_args()
    asyncio.run(ReportGeneratorV2(args.config).generate_report(args.step))

if __name__ == "__main__":
    main()